from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Tiêu đề
title = document.add_heading('KẾ HOẠCH TRIỂN KHAI HỆ THỐNG VIDEO RETRIEVAL AIC 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Báo cáo này trình bày chi tiết về kiến trúc hệ thống, lựa chọn công nghệ và lộ trình phát triển ứng dụng tìm kiếm video phục vụ cho cuộc thi HCMC AI Challenge (AIC) 2026.').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Lựa chọn Công nghệ (Tech Stack)
document.add_heading('1. LỰA CHỌN CÔNG NGHỆ (TECH STACK)', level=1)

document.add_heading('1.1. Backend (Xử lý Core & AI)', level=2)
document.add_paragraph('- Ngôn ngữ:', style='List Bullet')
document.add_paragraph('Python 3.10+ (Bắt buộc vì hệ sinh thái AI).')
document.add_paragraph('- Framework:', style='List Bullet')
document.add_paragraph('FastAPI (Cực kỳ nhanh, hỗ trợ xử lý bất đồng bộ, tự động sinh tài liệu API, lý tưởng để viết server AI).')
document.add_paragraph('- Vector Database:', style='List Bullet')
document.add_paragraph('FAISS (Facebook AI Similarity Search). FAISS chạy hoàn toàn trên RAM (In-memory), không cần cài đặt phức tạp như Docker, tốc độ tìm kiếm trên vài triệu vector siêu nhanh, lý tưởng cho máy tính cá nhân.')
document.add_paragraph('- AI Models:', style='List Bullet')
document.add_paragraph('  + Text Encoder: clip-ViT-B-32 (HuggingFace Transformers) để biến câu truy vấn thành vector.')
document.add_paragraph('  + Translation: Googletrans hoặc Gemini/OpenAI API để dịch Tiếng Việt sang Tiếng Anh trước khi đưa vào CLIP.')
document.add_paragraph('  + VQA (Vision-Language Model): API của Google Gemini 1.5 Flash hoặc OpenAI GPT-4o mini để xử lý câu hỏi nhanh gọn không đòi hỏi cấu hình GPU cục bộ.')

document.add_heading('1.2. Frontend (Giao diện người dùng)', level=2)
document.add_paragraph('- Framework:', style='List Bullet')
document.add_paragraph('ReactJS (Vite) + Vanilla CSS (hoặc TailwindCSS).')
document.add_paragraph('- Lý do:', style='List Bullet')
document.add_paragraph('Giao diện cần linh hoạt hiển thị Grid hàng trăm ảnh kết quả, click xem video xung quanh ảnh, và có chức năng "Giỏ hàng" lưu lại các kết quả Submit (giống Youtube). (Nếu đội không chuyên React, có thể chuyển sang dùng Streamlit của Python).')

# 2. Thiết kế Kiến trúc Hệ thống
document.add_heading('2. THIẾT KẾ KIẾN TRÚC HỆ THỐNG (SYSTEM ARCHITECTURE)', level=1)

document.add_heading('2.1. Giai đoạn 1: Chuẩn bị Dữ liệu (Offline / Tiền xử lý)', level=2)
document.add_paragraph('1. Script nạp dữ liệu: Đọc toàn bộ file .npy (CLIP features) từ thư mục BTC đưa vào một index của FAISS và lưu lại (VD: aic_features.index).', style='List Number')
document.add_paragraph('2. Cơ sở dữ liệu Map: Tạo SQLite hoặc JSON để ánh xạ ID của vector trong FAISS sang tên video_id và frame_id thực tế.', style='List Number')

document.add_heading('2.2. Giai đoạn 2: Máy chủ Tìm kiếm (Online / Lúc thi đấu)', level=2)
document.add_paragraph('Quy trình hoạt động (User Flow) của ứng dụng:', style='List Paragraph')
document.add_paragraph('1. Dạng Textual KIS:', style='List Number')
document.add_paragraph('Người dùng nhập câu truy vấn trên UI -> Gọi API Backend -> Dịch sang Tiếng Anh -> Mã hoá bằng CLIP text encoder -> Truy vấn vào FAISS -> Trả về UI danh sách Top 100 hình ảnh. Người dùng xác nhận, chọn ảnh và bấm "Add to Submission".')
document.add_paragraph('2. Dạng VQA:', style='List Number')
document.add_paragraph('Người dùng chọn 1 bức ảnh tìm được -> Gõ câu hỏi -> Backend gửi hình ảnh và câu hỏi lên API Gemini/GPT-4o -> Trả về câu trả lời -> "Add to Submission".')
document.add_paragraph('3. Export Dữ Liệu:', style='List Number')
document.add_paragraph('Cuối giờ thi, bấm "Export to Codabench", hệ thống tự động xuất các file .csv đúng định dạng và nén thành submission.zip để nộp.')

# 3. Lộ trình Phát triển
document.add_heading('3. LỘ TRÌNH PHÁT TRIỂN (ROADMAP)', level=1)
document.add_paragraph('BƯỚC 1:', style='List Bullet')
document.add_paragraph('Thiết lập môi trường Python, cài đặt FastAPI, FAISS, Transformers. Xây dựng Script Offline để lập chỉ mục (index) dữ liệu .npy.')
document.add_paragraph('BƯỚC 2:', style='List Bullet')
document.add_paragraph('Viết các API cho Backend (API tìm kiếm hình ảnh bằng Text, API xử lý VQA qua Gemini/OpenAI).')
document.add_paragraph('BƯỚC 3:', style='List Bullet')
document.add_paragraph('Xây dựng Giao diện Frontend bằng ReactJS hiển thị danh sách ảnh và tương tác người dùng.')
document.add_paragraph('BƯỚC 4:', style='List Bullet')
document.add_paragraph('Thiết kế chức năng Export ra file .csv theo chuẩn của nền tảng Codabench.')

# Lưu ý
document.add_heading('LƯU Ý DÀNH CHO ĐỘI THI', level=1)
document.add_paragraph('Để dự án có thể bắt tay vào viết mã nguồn ngay, đội thi cần chuẩn bị sẵn:')
document.add_paragraph('- Máy tính chạy hệ điều hành (Windows/Linux/Mac) có cài đặt Python 3.10+.', style='List Bullet')
document.add_paragraph('- Đảm bảo có token API của OpenAI hoặc Google Cloud (Gemini) để phục vụ cho tính năng hỏi đáp hình ảnh.', style='List Bullet')
document.add_paragraph('- Sắp xếp nhân sự phụ trách Frontend (React/HTML/CSS) và Backend (Python).', style='List Bullet')

document.save('Ke_Hoach_Trien_Khai_AIC_2026.docx')
