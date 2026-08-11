from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Tiêu đề
title = document.add_heading('BÁO CÁO TOÀN DIỆN: CHI TIẾT VÀ GIẢI PHÁP HỘI THI AIC 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Báo cáo này tổng hợp chi tiết nội dung truy vấn, phương pháp đánh giá, thông tin dữ liệu của vòng sơ tuyển AIC, đồng thời bổ sung hướng dẫn chi tiết cách thực hiện, thiết kế hệ thống AI và giải đáp các thắc mắc thường gặp.').alignment = WD_ALIGN_PARAGRAPH.CENTER

# PHẦN 1: THỂ LỆ VÀ ĐỀ BÀI
document.add_heading('PHẦN 1: TỔNG QUAN YÊU CẦU VÀ THỂ LỆ', level=1)

document.add_heading('1.1. Các dạng truy vấn', level=2)
document.add_paragraph('1. Textual KIS:', style='List Number')
document.add_paragraph('Tìm kiếm sự kiện dựa trên mô tả văn bản. Cần trả về video_id và frame_id.')
document.add_paragraph('2. Truy vấn Hỏi–Đáp (Q&A):', style='List Number')
document.add_paragraph('Tìm khoảnh khắc sự kiện và trả lời câu hỏi liên quan. Trả về video_id, frame_id, và answer (ngắn gọn, tối đa 100 ký tự).')
document.add_paragraph('3. TRAKE (Căn chỉnh chuỗi sự kiện thời gian):', style='List Number')
document.add_paragraph('Đòi hỏi tính chính xác cao để tìm ra một chuỗi các khung hình (semantic keyframes) diễn ra theo đúng thứ tự thời gian của sự kiện.')

document.add_heading('1.2. Phương pháp đánh giá và Dữ liệu', level=2)
document.add_paragraph('Đánh giá dựa trên Điểm Tương Quan (R-Score) từ 0 đến 1. Điểm cuối cùng là trung bình cộng của Top-k R-Score (k = 1, 5, 20, 50, 100).')
document.add_paragraph('Dữ liệu đợt 1 cung cấp:', style='List Paragraph')
document.add_paragraph('- Videos và Keyframes (Hình ảnh gốc).', style='List Bullet')
document.add_paragraph('- Objects (JSON): Tọa độ vật thể được trích xuất bằng Faster R-CNN.', style='List Bullet')
document.add_paragraph('- CLIP features (.npy): Đặc trưng không gian vector (clip-ViT-B-32) của tất cả khung hình.', style='List Bullet')

# PHẦN 2: CHIẾN LƯỢC VÀ GIẢI PHÁP AI
document.add_heading('PHẦN 2: PHÂN TÍCH GIẢI PHÁP VÀ HƯỚNG DẪN THỰC HIỆN', level=1)
document.add_paragraph('Trọng tâm của cuộc thi KHÔNG phải là tự huấn luyện (train) một mô hình AI từ đầu. BTC đã cung cấp sẵn đặc trưng (features) nén bởi siêu mô hình CLIP và Faster R-CNN. Đội thi cần "lắp ráp" và "tích hợp" các khối này thành một Hệ thống Tìm kiếm (Search Engine).')

document.add_heading('2.1. Giải pháp cho Textual KIS', level=2)
document.add_paragraph('- Kỹ thuật: Zero-shot Retrieval & Vector Search.', style='List Bullet')
document.add_paragraph('- Cách làm: Sử dụng Vector Database (FAISS, Milvus). Đưa toàn bộ file .npy của BTC vào database. Khi có câu hỏi, dùng mô hình Text Encoder của CLIP biến văn bản thành vector, sau đó dùng thuật toán tìm kiếm khoảng cách gần nhất (Cosine Similarity) để tìm ra khung hình khớp mô tả. Thời gian phản hồi chỉ dưới 1 giây.', style='List Bullet')

document.add_heading('2.2. Giải pháp cho VQA', level=2)
document.add_paragraph('- Kỹ thuật: Kết hợp Retrieval (Truy xuất) và VLM (Mô hình ngôn ngữ thị giác lớn).', style='List Bullet')
document.add_paragraph('- Cách làm: Dùng phương pháp ở KIS để tìm ra bức ảnh chứa cảnh diễn ra câu hỏi. Sau đó, truyền ảnh này cùng câu hỏi vào API của các AI như GPT-4o, Gemini 1.5 Pro hoặc Qwen-VL để AI trả lời. Với dạng câu hỏi đếm số lượng, nên viết code đọc trực tiếp số lượng bounding box trong file Objects JSON (Faster R-CNN) để có kết quả chính xác hơn so với việc bắt VLM đếm.', style='List Bullet')

document.add_heading('2.3. Giải pháp cho TRAKE', level=2)
document.add_paragraph('- Kỹ thuật: Temporal Alignment (Căn chỉnh thời gian).', style='List Bullet')
document.add_paragraph('- Cách làm: Chia nhỏ chuỗi hành động, tìm các khung hình liên quan cho từng hành động nhỏ. Sau đó sử dụng thuật toán Dynamic Time Warping (DTW) để ép các khung hình này phải nằm theo đúng trình tự thời gian (t1 < t2 < t3) trên cùng một video.', style='List Bullet')

document.add_heading('2.4. Kiến trúc Hệ thống đề xuất', level=2)
document.add_paragraph('Để hoàn thành bài thi trong 3 tiếng giới hạn, đội thi phải xây dựng một phần mềm/ứng dụng hoàn chỉnh:', style='List Paragraph')
document.add_paragraph('- Backend (Python/FastAPI): Tiếp nhận truy vấn, xử lý vector qua FAISS, và kết nối với các AI LLM.', style='List Bullet')
document.add_paragraph('- Frontend (Web UI): Giao diện tìm kiếm nội bộ (giống Google/Youtube) để thành viên đội nhập câu hỏi, xem trước video kết quả (cực kỳ quan trọng để mắt người duyệt lại độ chính xác).', style='List Bullet')
document.add_paragraph('- Module Export: Tự động xuất kết quả thành định dạng CSV chuẩn và nén ZIP (thư mục submission/) để upload lên hệ thống Codabench.', style='List Bullet')

# PHẦN 3: HỎI ĐÁP (Q&A)
document.add_heading('PHẦN 3: HỎI ĐÁP VÀ GIẢI ĐÁP THẮC MẮC (Q&A)', level=1)

document.add_heading('Câu hỏi 1: Đội thi có cần xây dựng và huấn luyện mô hình AI mới hoàn toàn không?', level=3)
document.add_paragraph('Trả lời: KHÔNG. Cuộc thi đánh giá năng lực "Kỹ sư hệ thống" (System Engineering). Bạn tận dụng đặc trưng AI đã được BTC trích xuất (CLIP, Faster R-CNN) để xây dựng hệ thống truy xuất (Retrieval System). Bạn không cần GPU khủng để train AI trong cuộc thi, mà cần kỹ năng lập trình ghép nối API và Database.')

document.add_heading('Câu hỏi 2: Có được phép sử dụng các AI như ChatGPT hay Gemini trong lúc thi không?', level=3)
document.add_paragraph('Trả lời: ĐƯỢC PHÉP. Đội thi tự túc về mặt phần cứng và phần mềm, việc sử dụng các API thương mại (như OpenAI, Gemini) hoặc các mô hình mã nguồn mở để giải bài (đặc biệt là dạng VQA) là hướng đi được khuyến khích để nâng cao R-Score.')

document.add_heading('Câu hỏi 3: Nếu đáp án văn bản của dạng VQA có dấu phẩy hoặc ngoặc kép thì làm sao?', level=3)
document.add_paragraph('Trả lời: File kết quả CSV rất nhạy cảm với các ký tự này. Đội thi phải escape (thoát) chuỗi cẩn thận bằng cách bọc toàn bộ chuỗi đáp án bằng dấu ngoặc kép "...". Các dấu ngoặc kép bên trong thì phải chuyển thành hai dấu ngoặc kép liên tiếp "".')

document.save('Bao_Cao_Toan_Dien_AIC_2026.docx')
