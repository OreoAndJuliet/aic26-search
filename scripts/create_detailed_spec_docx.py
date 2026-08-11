import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title (Simple and Short)
title = doc.add_heading('ĐẶC TẢ HỆ THỐNG AIC 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Mục đích tài liệu:', style='Heading 2')
doc.add_paragraph('Cung cấp đặc tả hệ thống và giao ước kỹ thuật cho hai nhóm Frontend và Backend. Đảm bảo hệ thống Tìm kiếm Video (Interactive Video Retrieval) hoạt động trơn tru trong thời gian thi đấu.')

# PHẦN 1
doc.add_heading('PHẦN 1: KIẾN TRÚC HỆ THỐNG TỔNG THỂ', level=1)
doc.add_paragraph('Bài toán của AIC 2026 yêu cầu tìm kiếm chính xác khung hình và video (Textual KIS, VQA, TRAKE) với tốc độ cao. Hệ thống áp dụng mô hình "Human-in-the-loop" (Con người kiểm duyệt kết quả của máy).')

doc.add_heading('1.1. Kiến trúc High-Level', level=2)
doc.add_paragraph('- Frontend (Web App): Giao diện tương tác, hiển thị kết quả trực quan (ảnh grid, video player), hỗ trợ duyệt và nộp bài.', style='List Bullet')
doc.add_paragraph('- Backend (API Server): Xử lý logic, kết nối AI Translator (Dịch thuật), AI Models (CLIP, VLM) và Vector Database (FAISS/Milvus).', style='List Bullet')

# PHẦN 2
doc.add_heading('PHẦN 2: ĐẶC TẢ CHI TIẾT NHÓM FRONTEND', level=1)

doc.add_heading('2.1. Yêu cầu Giao diện', level=2)
doc.add_paragraph('1. Màn hình Tìm kiếm: Thanh tìm kiếm, Dropdown chọn loại truy vấn (KIS, VQA, TRAKE).', style='List Number')
doc.add_paragraph('2. Lưới hiển thị (Grid): Hiển thị ảnh kèm R-Score, video_id, frame_id.', style='List Number')
doc.add_paragraph('3. Video Player: Modal phát đoạn video tương ứng với khung hình.', style='List Number')
doc.add_paragraph('4. Giỏ hàng (Cart): Lưu tạm đáp án và xuất CSV/ZIP.', style='List Number')

doc.add_heading('2.2. Yêu cầu Kỹ thuật', level=2)
doc.add_paragraph('- Áp dụng "Lazy Loading" cho Grid ảnh để tránh tràn RAM.', style='List Bullet')
doc.add_paragraph('- Công nghệ tự do, khuyến nghị React + Vite.', style='List Bullet')

# PHẦN 3
doc.add_heading('PHẦN 3: ĐẶC TẢ CHI TIẾT NHÓM BACKEND & AI', level=1)

doc.add_heading('3.1. Quy trình xử lý dữ liệu', level=2)
doc.add_paragraph('- Pre-processing: Nạp dữ liệu CLIP (.npy), Objects (.json) vào Vector DB.', style='List Bullet')
doc.add_paragraph('- Hosting: Host static server cho ảnh (.jpg) và video (.mp4).', style='List Bullet')

doc.add_heading('3.2. Động cơ Xử lý (AI Engines)', level=2)
doc.add_paragraph('- KIS Engine: Kết nối AI Translator dịch câu tiếng Việt sang tiếng Anh. Dùng CLIP text-encoder tạo vector, tìm kiếm trên Vector DB.', style='List Bullet')
doc.add_paragraph('- VQA Engine: Lấy Top 5 ảnh từ KIS, gọi AI Vision (LLM/VLM) để sinh câu trả lời.', style='List Bullet')
doc.add_paragraph('- TRAKE Engine: Chạy DTW để căn chỉnh thời gian cho chuỗi sự kiện.', style='List Bullet')

# PHẦN 4
doc.add_heading('PHẦN 4: GIAO ƯỚC API (CONTRACTS)', level=1)

doc.add_heading('4.1. Static Files', level=2)
doc.add_paragraph('Ảnh: http://<backend_url>/static/keyframes/{video_name}/{frame_id}.jpg')
doc.add_paragraph('Video: http://<backend_url>/static/videos/{video_name}.mp4')

doc.add_heading('4.2. API /search', level=2)
doc.add_paragraph('Endpoint: POST /api/v1/search')
doc.add_paragraph('Request:\n{\n  "type": "KIS",\n  "text": "Một người...",\n  "question": null,\n  "top_k": 50\n}')
doc.add_paragraph('Response:\n{\n  "status": "success",\n  "results": [\n    {\n      "video_id": "L01_V001",\n      "frame_id": 1500,\n      "thumbnail_url": "...",\n      "answer": null\n    }\n  ]\n}')

doc.add_heading('4.3. Xuất File', level=2)
doc.add_paragraph('- Không header. Format: <video_id>, <frame_id>, "<answer>"')

# PHẦN 5
doc.add_heading('PHẦN 5: LỘ TRÌNH (ROADMAP)', level=1)
doc.add_paragraph('1. Giai đoạn 1 (Tuần 1-2):', style='List Number')
doc.add_paragraph('- Backend: Xây dựng API chuẩn, kết nối AI Translator (Google Translate API hoặc model nội bộ) và kết nối AI CLIP để tìm kiếm ảnh cơ bản. Phục vụ ảnh/video tĩnh.', style='List Bullet')
doc.add_paragraph('- Frontend: Thiết kế giao diện (UI) và gọi API thật/mock từ Backend.', style='List Bullet')
doc.add_paragraph('2. Giai đoạn 2 (Tuần 3-4):', style='List Number')
doc.add_paragraph('- Backend: Hoàn thiện VQA (tích hợp VLM) và TRAKE (thuật toán DTW).', style='List Bullet')
doc.add_paragraph('- Frontend: Làm chức năng xuất CSV/ZIP. Sau đó ghép team hỗ trợ Backend.', style='List Bullet')
doc.add_paragraph('3. Giai đoạn 3 (Thực chiến): Thi thử nội bộ, tối ưu hiệu năng.', style='List Number')

doc.save('Dac_Ta_AIC_2026.docx')
