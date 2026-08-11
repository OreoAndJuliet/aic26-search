from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Tiêu đề
title = document.add_heading('BÁO CÁO TỔNG QUAN: HỘI THI THỬ THÁCH TRÍ TUỆ NHÂN TẠO THÀNH PHỐ HỒ CHÍ MINH (AIC)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Báo cáo này tổng hợp chi tiết và khoa học về nội dung truy vấn, phương pháp đánh giá và thông tin dữ liệu của vòng sơ tuyển AIC, dựa trên Đề bài chính thức và nền tảng Codabench.').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. NỘI DUNG CÁC TRUY VẤN VÒNG SƠ TUYỂN
document.add_heading('1. NỘI DUNG CÁC TRUY VẤN VÒNG SƠ TUYỂN', level=1)

document.add_heading('1.1. Truy vấn dạng 1: Tìm kiếm chính xác theo văn bản (Textual KIS)', level=2)
document.add_paragraph('Đây là nhiệm vụ tìm kiếm sự kiện dựa trên mô tả bằng văn bản. Ban giám khảo cung cấp mô tả ngôn ngữ tự nhiên. Các đội cần định vị đoạn video chứa sự kiện bằng cách chỉ ra một khung hình bất kỳ thuộc đoạn video đó.')
document.add_paragraph('Ví dụ: Truy vấn "Tìm video về một diễn giả mặc áo đỏ phát biểu tại một cuộc họp báo ngoài trời, phía sau có nhiều cây xanh." — Kết quả nộp: video_id = video_abc.mp4, frame_id = 1500.', style='Intense Quote')

document.add_heading('1.2. Truy vấn dạng 2: Truy vấn dạng Hỏi–Đáp (Q&A)', level=2)
document.add_paragraph('Nhiệm vụ tìm kiếm sự kiện và trích xuất thông tin cụ thể từ video. Ban giám khảo cung cấp mô tả sự kiện và câu hỏi. Đội thi cần tìm ra chính xác khoảnh khắc liên quan và trả lời câu hỏi bằng tiếng Việt hoặc Anh.')
document.add_paragraph('Ví dụ: Truy vấn "Trong video về lễ trao giải thưởng âm nhạc, có bao nhiêu người lên sân khấu để nhận giải thưởng lớn nhất?" — Kết quả nộp: video_id = video_xyz.mp4, frame_id = 3450, answer = "5".', style='Intense Quote')

document.add_heading('1.3. Truy vấn dạng 3: Truy xuất và căn chỉnh sự kiện video theo thời gian (TRAKE)', level=2)
document.add_paragraph('Nhiệm vụ phức hợp đòi hỏi độ chính xác cao trong cả truy xuất và căn chỉnh thời gian khoảnh khắc. Chia thành 2 giai đoạn:')
document.add_paragraph(' - Giai đoạn 1 (Truy xuất): Tìm ra 1 video duy nhất chứa chuỗi sự kiện khớp nhất.', style='List Bullet')
document.add_paragraph(' - Giai đoạn 2 (Căn chỉnh): Xác định 1 khung hình ngữ nghĩa (semantic keyframe) duy nhất cho mỗi giai đoạn của chuỗi sự kiện.', style='List Bullet')

# 2. PHƯƠNG PHÁP ĐÁNH GIÁ
document.add_heading('2. PHƯƠNG PHÁP ĐÁNH GIÁ VÒNG SƠ TUYỂN', level=1)
document.add_paragraph('Mỗi truy vấn gửi tối đa 100 câu trả lời. Mỗi câu được chấm Điểm Tương Quan (R-Score) từ 0 đến 1.')

document.add_heading('2.1. Điểm Tương Quan (R-Score)', level=2)
document.add_paragraph('2.1.1. Textual KIS:', style='List Number')
document.add_paragraph('Định dạng trả lời: <video_id>, <frame_id>')
document.add_paragraph('Điều kiện: Đúng video và frame_id nằm trong khoảng đáp án [s, e].')

document.add_paragraph('2.1.2. Q&A:', style='List Number')
document.add_paragraph('Định dạng trả lời: <video_id>, <frame_id>, <answer>')
document.add_paragraph('Điều kiện: Đúng video, đúng khoảng frame, và answer đúng ngữ nghĩa.')

document.add_paragraph('2.1.3. TRAKE:', style='List Number')
document.add_paragraph('Định dạng trả lời: <video_id>, <frame_id_1>, ..., <frame_id_n>')
document.add_paragraph('Điều kiện: Sai video -> 0 điểm. Đúng video, tính tỉ lệ số khung hình khớp với các khoảng đáp án của chuỗi sự kiện.')

document.add_heading('2.2. Điểm Cuối Cùng (Final Score)', level=2)
document.add_paragraph('Sử dụng Top-k R-Score (R@k) với k thuộc {1, 5, 20, 50, 100}. Điểm cuối cùng là trung bình cộng của 5 giá trị R@k.')

# 3. THÔNG TIN DỮ LIỆU
document.add_heading('3. THÔNG TIN DỮ LIỆU VÒNG SƠ TUYỂN – ĐỢT 1', level=1)
document.add_paragraph('Dữ liệu bao gồm các thành phần:', style='List Paragraph')
document.add_paragraph('- Videos: Chứa video gốc.', style='List Bullet')
document.add_paragraph('- Keyframes: Các keyframe trích xuất từ video.', style='List Bullet')
document.add_paragraph('- Objects: File JSON chứa vật thể từ mô hình Faster R-CNN (OpenImages V4).', style='List Bullet')
document.add_paragraph('- CLIP features: NPY file chứa đặc trưng clip-ViT-B-32.', style='List Bullet')
document.add_paragraph('- Metadata: Thông tin gốc từ YouTube.', style='List Bullet')
document.add_paragraph('Link tải dữ liệu đợt 1: https://docs.google.com/spreadsheets/d/1rfn1fieTThS_Ki3SIoJ6uXOx2AhMq7wGCak6W4jZyZM/edit?usp=sharing')
document.add_paragraph('Lưu ý: Dữ liệu chính thức là Video, các thành phần khác chỉ để hỗ trợ. Đợt 2 sẽ có thêm dữ liệu bổ sung.')

# 4. TIMELINE
document.add_heading('4. LỘ TRÌNH VÀ CÁC QUY ĐỊNH NỘP BÀI (CODABENCH)', level=1)
table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Giai đoạn'
hdr_cells[1].text = 'Thời gian'
hdr_cells[2].text = 'Mô tả'

phases = [
    ('Nộp thử nghiệm', '24/08 - 30/08/2025', 'Test hệ thống, lỗi định dạng'),
    ('Lượt 1 (Round 1)', '31/08/2025 (09:00 - 11:59)', 'Đề: AIC25-Pack1-GroupA'),
    ('Lượt 2 (Round 2)', '07/09/2025 (09:00 - 11:59)', 'Đề: AIC25-Pack2-GroupA'),
    ('Lượt 3 (Round 3)', '14/09/2025 (09:00 - 11:59)', 'Đề: AIC25-Pack3-GroupA')
]
for p, t, d in phases:
    row_cells = table.add_row().cells
    row_cells[0].text = p
    row_cells[1].text = t
    row_cells[2].text = d

document.add_paragraph('\nĐịnh dạng nộp bài (Codabench): Tất cả kết quả lưu dạng .csv không header, để trong thư mục "submission/", sau đó nén lại thành file .zip để tải lên.')

document.save('Bao_Cao_Chi_Tiet_AIC.docx')
